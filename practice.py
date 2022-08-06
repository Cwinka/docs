import argparse
import sys
from pathlib import Path
from docx import Document
from docx.document import Document as HintDocument
from typing import Callable
from loguru import logger
from interfaces import UsurtDoc, XlsxData, UsurtBaseTable, XlsxDataParser, UnsetFieldsError
from docx.shared import Inches


class UsurtXlsxPracticeData(XlsxData):
    aim: str
    aim_kind: str
    grade: str
    faculty: str
    group: str
    study_type: str
    specialization: str
    year_period: str
    date_period: str
    director: str
    director_name: str
    orgs: list[list[tuple[str, str | None]]]

    def __init__(self):
        self.orgs = []
        self._orgs_key = 'Группа организаций. [Имя организации]. ФИО студентов. Группа, форма обучения'
        self.columns: dict[str, tuple[[Callable[[str], None]], int, int]] = {
            "Вид практики": (self._set_it('aim'), 1, 1),
            "Тип практики": (self._set_it('aim_kind'), 1, 1),
            "Курс": (self._set_it('grade'), 1, 1),
            "Факультет": (self._set_it('faculty'), 1, 1),
            "Группа": (self._set_it('group'), 1, 1),
            "Тип обучения": (self._set_it('study_type'), 1, 1),
            "Специализация": (self._set_it('specialization'), 1, 1),
            "Период практики (годы)": (self._set_it('year_period'), 1, 1),
            "Период практики (дни)": (self._set_it('date_period'), 1, 1),
            "Должность руководителя практики": (self._set_it('director'), 1, 1),
            "ФИО руководителя практики": (self._set_it('director_name'), 1, 1)
        }

    def get_setter(self, xlsx_field: str) -> tuple[[Callable[[str], None]], int | None, int] | None:
        try:
            if xlsx_field == self._orgs_key:
                # поля организация
                return self._set_in_list(self.orgs), None, 2
            return self.columns.pop(xlsx_field)
        except KeyError:
            return

    def get_unset_fields(self) -> tuple[str]:
        return tuple(self.columns)


class UsurtPracticeDoc(UsurtDoc):

    def __init__(self, data: UsurtXlsxPracticeData):
        self.data = data

    def intro(self):
        yield 'ФЕДЕРАЛЬНОЕ АГЕНТСТВО ЖЕЛЕЗНОДОРОЖНОГО ТРАНСПОРТА'
        yield 'Федеральное государственное бюджетное образовательное учреждение высшего образования'
        yield '«Уральский государственный университет путей сообщения»'
        yield '(ФГБОУ ВО УрГУПС)'

        through_spaces = ('П Р И К А З', '«___»____________ 2022 г. №___________', 'Екатеринбург')
        for p in through_spaces:
            yield ''
            yield p

        morpy = self._morf_to(self.data.aim, 'loct')  # предложный
        pre = 'Об' if morpy[0].lower() in 'ауеыоэию' else 'О'
        yield f'{pre} {morpy}'
        yield f'обучающихся {self.data.grade} курса {self.data.faculty}'

    def order(self):
        aim = f'{self.data.aim} ({self.data.aim_kind})'
        yield f'В соответствии с календарным учебным графиком на {self.data.year_period} учебный год'
        yield f'ПРИКАЗЫВАЮ'
        study_type = self._morf_to(self.data.study_type, "nomn")[:-2] + 'й'
        yield f'Направить нижепоименованных студентов {self.data.grade} курса ' \
              f'{study_type} формы обучения ' \
              f'направления подготовки {self.data.specialization} для прохождения ' \
              f'{self._morf_to(aim, "gent")} в период с ' \
              f'{self.data.date_period} на следующие объекты и ' \
              'утвердить руководителей практики УрГУПС:'

    def make_table(self, doc: HintDocument) -> 'UsurtBaseTable':
        t = UsurtBaseTable(doc, columns_width=(Inches(5),))
        t.add_row('Фамилия Имя Отчество обучающегося', 'Группа, форма обучения (ц, б, к)',
                  'Руководитель практики от УрГУПС')
        t.merge((0,), cells=(2, 3))
        t.add_row('', '', 'Должность', 'Ф.И.О.')
        t.merge((0, 1), cells=(0, 1))
        return t

    def make_all_tables(self, doc: 'HintDocument'):
        for org in self.data.orgs:
            org_common_name, _ = org[0]
            assert _ is None, f'Неверный формат ввода группы организаций. Поле справа от ' \
                              f'группы организации должно быть пустым.'
            doc.add_paragraph('')
            doc.add_paragraph(org_common_name)
            students = self.make_table(doc)
            for part in org[1:]:
                # конкретное имя организации в группе организацй
                if part[1] is None:
                    row = students.add_row(part[0], char_style='Strong')
                    students.merge((row,), cells=(0, 3))
                else:
                    students.add_row(part[0], part[1], self.data.director, self.data.director_name)
        doc.add_paragraph('')

    def ending(self):
        yield f'Общая ответственность за проведение {self._morf_to(self.data.aim, "gent")} ' \
              f'возлагается на заведующего кафедрой  «Мехатроника», к.ф.-м.н. В. С. Тарасяна.'
        yield 'Проректор по учебной работе'
        yield 'и связям с производством						Н. Ф. Сирина'


def main():
    logger.remove()
    logger.add(sys.stdout, colorize=True, format="<level>{level}</level> | <level>{message}</level>")
    parser = argparse.ArgumentParser()
    parser.add_argument('xlsx', type=str, help='путь до шаблона xlsx документа.')
    parser.add_argument('-o', '-out', type=str, help='путь до нового doc документа, по умолчанию имя документа и '
                                                     'директория совпадают с xlsx путём.')
    args = parser.parse_args()
    xlsx_path = Path(args.xlsx)
    out = Path(args.o) if args.o else Path(f'{xlsx_path.stem}.docx')

    doc: HintDocument = Document()
    try:
        info = UsurtPracticeDoc(XlsxDataParser.parse(xlsx_path, UsurtXlsxPracticeData))
    except UnsetFieldsError as e:
        logger.error(e.err)
        exit(1)
    else:
        for part in (info.intro(), info.order()):
            for text in part:
                doc.add_paragraph(text)

        info.make_all_tables(doc)

        for text in info.ending():
            doc.add_paragraph(text)
        logger.info(f'Документ {out.as_posix()} успешно создан по шаблону {xlsx_path.as_posix()}.')
        doc.save(out)


if __name__ == '__main__':
    main()
